"""
Document loader — raw file → text + metadata.

Why pymupdf (fitz) over pypdf or pdfplumber?
- pymupdf is the fastest pure-Python PDF parser by a large margin.
- It preserves reading order better than pypdf on multi-column layouts,
  which matters for contracts with signature blocks and tables.
- pdfplumber is better for table extraction but we don't need that here —
  we're extracting clause text, not structured table data.

Why extract page-by-page and track page numbers?
source_chunk_id in ExtractedClause needs to trace back to a specific page
for citations. If we concatenate the whole document first, we lose that
mapping. Page-level metadata is attached at load time so it flows through
chunking and indexing without extra bookkeeping.

Why python-docx for DOCX and not converting to PDF first?
Converting DOCX → PDF introduces layout artifacts and requires an external
tool (LibreOffice or similar). python-docx reads the XML directly and gives
us clean paragraph-level text. Simpler, faster, no external dependency.
"""

from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

import fitz  # pymupdf
from docx import Document as DocxDocument


@dataclass
class PagedText:
    """Raw text for a single page/section, with its origin metadata."""
    doc_id: str
    file_path: str
    page_number: int          # 1-indexed
    text: str
    total_pages: int


@dataclass
class LoadedDocument:
    """All pages from a single document after loading."""
    doc_id: str
    file_path: str
    total_pages: int
    pages: list[PagedText] = field(default_factory=list)

    @property
    def full_text(self) -> str:
        return "\n\n".join(p.text for p in self.pages)


def load_document(file_path: str, doc_id: Optional[str] = None) -> LoadedDocument:
    """
    Load a PDF or DOCX file and return page-level text with metadata.
    doc_id defaults to the file stem if not provided.
    """
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"Document not found: {file_path}")

    resolved_id = doc_id or path.stem
    suffix = path.suffix.lower()

    if suffix == ".pdf":
        return _load_pdf(str(path), resolved_id)
    elif suffix in (".docx", ".doc"):
        return _load_docx(str(path), resolved_id)
    elif suffix == ".txt":
        return _load_txt(str(path), resolved_id)
    else:
        raise ValueError(f"Unsupported file type: {suffix}. Expected .pdf, .docx, or .txt")


def _load_pdf(file_path: str, doc_id: str) -> LoadedDocument:
    doc = fitz.open(file_path)
    total_pages = len(doc)
    pages: list[PagedText] = []

    for page_num in range(total_pages):
        page = doc[page_num]
        # "text" layout preserves reading order better than raw extraction
        text = str(page.get_text("text")).strip()
        if not text:
            continue  # skip blank/image-only pages
        pages.append(PagedText(
            doc_id=doc_id,
            file_path=file_path,
            page_number=page_num + 1,
            text=text,
            total_pages=total_pages,
        ))

    doc.close()
    return LoadedDocument(
        doc_id=doc_id,
        file_path=file_path,
        total_pages=total_pages,
        pages=pages,
    )


def _load_txt(file_path: str, doc_id: str) -> LoadedDocument:
    text = Path(file_path).read_text(encoding="utf-8", errors="replace")
    lines = [l for l in text.splitlines() if l.strip()]
    page_size = 50
    synthetic_pages = [lines[i : i + page_size] for i in range(0, len(lines), page_size)]

    pages: list[PagedText] = []
    for i, group in enumerate(synthetic_pages):
        pages.append(PagedText(
            doc_id=doc_id,
            file_path=file_path,
            page_number=i + 1,
            text="\n".join(group),
            total_pages=len(synthetic_pages),
        ))

    return LoadedDocument(
        doc_id=doc_id,
        file_path=file_path,
        total_pages=len(synthetic_pages),
        pages=pages,
    )


def _load_docx(file_path: str, doc_id: str) -> LoadedDocument:
    doc = DocxDocument(file_path)

    # DOCX has no concept of "pages" — we use paragraphs as the unit.
    # Group into synthetic ~50-paragraph "pages" so the metadata structure
    # stays consistent with PDF output downstream.
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    page_size = 50
    synthetic_pages = [
        paragraphs[i : i + page_size]
        for i in range(0, len(paragraphs), page_size)
    ]

    pages: list[PagedText] = []
    for i, group in enumerate(synthetic_pages):
        pages.append(PagedText(
            doc_id=doc_id,
            file_path=file_path,
            page_number=i + 1,
            text="\n".join(group),
            total_pages=len(synthetic_pages),
        ))

    return LoadedDocument(
        doc_id=doc_id,
        file_path=file_path,
        total_pages=len(synthetic_pages),
        pages=pages,
    )
